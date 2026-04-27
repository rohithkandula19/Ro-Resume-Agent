"""Export resumes to DOCX and PDF in any font from the library.

Input: a structured resume dict with keys:
  name, contact, summary, skills, experience, projects, education, certifications
Experience/projects entries: {title, org, dates, bullets: [str]}
Education entries: {degree, school, dates, notes}

Two output flavors:
  - styled: uses the user's chosen font + template style
  - ats_safe: Calibri/Arial, single column, minimal formatting
"""

import io
import os
import re
from typing import Any

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

from library.fonts import FONTS, fallback_for_ats


# ------------------------- TEMPLATE STYLE REGISTRY -------------------------

# Layout + accent color per template id. Layouts: minimal, modern-bar, serif-classic, sidebar, banner, structured
TEMPLATE_STYLE: dict[str, dict] = {
    "google_xyz":          {"layout": "minimal",       "accent": "#1a73e8"},
    "meta_engineering":    {"layout": "minimal",       "accent": "#1877f2"},
    "amazon_lp":           {"layout": "modern-bar",    "accent": "#ff9900"},
    "apple_minimal":       {"layout": "minimal",       "accent": "#111111"},
    "microsoft_modern":    {"layout": "modern-bar",    "accent": "#0078d4"},
    "shopify_commerce":    {"layout": "banner",        "accent": "#008060"},
    "fresher_student":     {"layout": "minimal",       "accent": "#7c3aed"},
    "netflix_senior_ic":   {"layout": "modern-bar",    "accent": "#e50914"},
    "us_standard":         {"layout": "minimal",       "accent": "#222222"},
    "international_global":{"layout": "minimal",       "accent": "#0d9488"},
    "swe_backend":         {"layout": "modern-bar",    "accent": "#4f46e5"},
    "devops_sre":          {"layout": "modern-bar",    "accent": "#0ea5e9"},
    "stripe_eng":          {"layout": "banner",        "accent": "#635bff"},
    "uber_pm":             {"layout": "banner",        "accent": "#000000"},
    "linkedin_recruiter":  {"layout": "minimal",       "accent": "#0a66c2"},
    "mckinsey_consulting": {"layout": "serif-classic", "accent": "#003366"},
    "bcg_bain":            {"layout": "serif-classic", "accent": "#00563f"},
    "goldman_ib":          {"layout": "serif-classic", "accent": "#b45309"},
    "big4":                {"layout": "serif-classic", "accent": "#c8102e"},
    "europass":            {"layout": "structured",    "accent": "#003399"},
    "uk_cv":               {"layout": "minimal",       "accent": "#1f2937"},
    "dach_lebenslauf":     {"layout": "serif-classic", "accent": "#374151"},
    "india_standard":      {"layout": "sidebar",       "accent": "#ea580c"},
    "middle_east":         {"layout": "minimal",       "accent": "#0f766e"},
    "balanced_hybrid":     {"layout": "modern-bar",    "accent": "#0d9488"},
    "modern_classic":      {"layout": "serif-classic", "accent": "#334155"},
    "academic_cv":         {"layout": "serif-classic", "accent": "#7f1d1d"},
    "career_change":       {"layout": "sidebar",       "accent": "#9333ea"},
    "frontend_ui":         {"layout": "sidebar",       "accent": "#ec4899"},
    "data_ml":             {"layout": "modern-bar",    "accent": "#2563eb"},
    "ro_signature":        {"layout": "shaded-header", "accent": "#1a1a1a"},
    # City Series
    "dublin_classic":      {"layout": "serif-classic", "accent": "#0d9488"},
    "sydney_modern":       {"layout": "sidebar",       "accent": "#1e3a5f"},
    "paris_elegant":       {"layout": "serif-classic", "accent": "#7c2d44"},
    "chicago_bold":        {"layout": "modern-bar",    "accent": "#1e293b"},
    "toronto_professional":{"layout": "modern-bar",    "accent": "#3730a3"},
    "milan_executive":     {"layout": "serif-classic", "accent": "#78716c"},
    # Specialized Industry
    "federal_gov":         {"layout": "structured",    "accent": "#1e3a5f"},
    "healthcare_clinical": {"layout": "minimal",       "accent": "#0891b2"},
    "educator_teacher":    {"layout": "serif-classic", "accent": "#15803d"},
}


def _template_style(template_name_or_id: str) -> dict:
    """Resolve by id, else by name match against the catalog, else fall back to minimal/black."""
    t = (template_name_or_id or "").strip()
    if t in TEMPLATE_STYLE:
        return TEMPLATE_STYLE[t]
    try:
        from library.templates import TEMPLATES
        for tpl in TEMPLATES:
            if tpl.get("name") == t or tpl.get("id") == t:
                return TEMPLATE_STYLE.get(tpl["id"], {"layout": "minimal", "accent": "#111111"})
    except Exception:
        pass
    return {"layout": "minimal", "accent": "#111111"}


# ------------------------- DOCX EXPORT -------------------------

def _docx_heading(doc: Document, text: str, font_name: str, size: int = 12,
                   bold: bool = True, color=(0x11, 0x11, 0x11)):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    return p


def _docx_text(doc: Document, text: str, font_name: str, size: int = 10,
                bold: bool = False, italic: bool = False, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def _docx_bullet(doc: Document, text: str, font_name: str, size: int = 10):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"• {text}")
    run.font.name = font_name
    run.font.size = Pt(size)
    return p


def build_docx(resume: dict[str, Any], font_name: str, font_size: float = 10.0) -> bytes:
    doc = Document()
    body_sz = int(max(8, min(13, round(font_size or 10))))
    name_sz = body_sz + 8
    head_sz = body_sz + 2

    # narrow margins for density
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Header
    name = resume.get("name") or "Your Name"
    contact = resume.get("contact") or ""
    _docx_text(doc, name, font_name, size=name_sz, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if contact:
        _docx_text(doc, contact, font_name, size=body_sz, align=WD_ALIGN_PARAGRAPH.CENTER)

    if resume.get("summary"):
        _docx_heading(doc, "Summary", font_name, size=head_sz)
        _docx_text(doc, resume["summary"], font_name, size=body_sz)

    if resume.get("skills"):
        _docx_heading(doc, "Skills", font_name, size=head_sz)
        skills = resume["skills"]
        if isinstance(skills, list):
            _docx_text(doc, " • ".join(skills), font_name, size=body_sz)
        else:
            _docx_text(doc, str(skills), font_name, size=body_sz)

    def section(title: str, items: list, include_bullets: bool = True):
        if not items:
            return
        _docx_heading(doc, title, font_name, size=head_sz)
        for item in items:
            header_parts = []
            if item.get("title"): header_parts.append(item["title"])
            if item.get("org"):   header_parts.append(item["org"])
            header = " — ".join(header_parts)
            dates = item.get("dates", "")
            line = f"{header}   {dates}" if dates else header
            _docx_text(doc, line, font_name, size=body_sz, bold=True)
            if include_bullets:
                for b in item.get("bullets", []) or []:
                    _docx_bullet(doc, b, font_name, size=body_sz)

    section("Experience", resume.get("experience") or [])
    section("Projects", resume.get("projects") or [])

    if resume.get("education"):
        _docx_heading(doc, "Education", font_name, size=head_sz)
        for ed in resume["education"]:
            degree = str(ed.get("degree") or "")
            school = str(ed.get("school") or "")
            dates  = ed.get("dates", "")
            if isinstance(dates, list): dates = " ".join(str(d) for d in dates)
            tail_bits = [x for x in (school, str(dates or "")) if x]
            tail = " — ".join(tail_bits)
            p = doc.add_paragraph()
            if degree:
                r = p.add_run(degree)
                r.font.name = font_name
                r.font.size = Pt(body_sz)
                r.font.bold = True
            if tail:
                r2 = p.add_run((" — " if degree else "") + tail)
                r2.font.name = font_name
                r2.font.size = Pt(body_sz)
                r2.font.bold = False
            if ed.get("notes"):
                _docx_text(doc, ed["notes"], font_name, size=body_sz)

    if resume.get("certifications"):
        _docx_heading(doc, "Certifications", font_name, size=head_sz)
        for c in resume["certifications"]:
            _docx_bullet(doc, c if isinstance(c, str) else c.get("name", str(c)),
                         font_name, size=body_sz)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ------------------------- PDF EXPORT -------------------------

_REGISTERED_FONTS: set[str] = set()
_BUNDLED_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "fonts")

_SEARCH_DIRS = [
    _BUNDLED_DIR,
    "/Library/Fonts",
    "/System/Library/Fonts",
    "/System/Library/Fonts/Supplemental",
    os.path.expanduser("~/Library/Fonts"),
    "/usr/share/fonts",
    "/usr/local/share/fonts",
]


def _find_ttf(basenames: set[str]) -> str | None:
    """First file whose basename (lowercased, no ext) is in the provided set."""
    exts = {".ttf", ".otf"}
    lowered = {b.lower() for b in basenames}
    for d in _SEARCH_DIRS:
        if not os.path.isdir(d):
            continue
        for root, _, files in os.walk(d):
            for f in files:
                base, ext = os.path.splitext(f)
                if ext.lower() in exts and base.lower() in lowered:
                    return os.path.join(root, f)
    return None


def _register_pdf_font(font_name: str) -> str:
    """Register a font family (regular + bold + italic) for use in PDF output.

    Bold and italic tags in ReportLab paragraphs require a full family, else they
    fall back to Helvetica. Bundled TTFs (assets/fonts/) take priority over system
    fonts. Returns the family name to pass as fontName, or 'Helvetica' on failure.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily

    if font_name in _REGISTERED_FONTS:
        return font_name
    if font_name == "Helvetica":
        return "Helvetica"

    squashed = font_name.replace(" ", "")

    def names_for(suffix: str) -> set[str]:
        s = suffix
        base = [font_name, squashed]
        if not s:
            return {n for n in base} | {f"{n}-regular" for n in base} | {f"{n}regular" for n in base}
        return {f"{n}-{s}" for n in base} | {f"{n}{s}" for n in base}

    regular_path = _find_ttf(names_for(""))
    if not regular_path:
        return "Helvetica"

    bold_path = _find_ttf(names_for("bold"))
    italic_path = _find_ttf(names_for("italic"))
    bolditalic_path = _find_ttf(names_for("bolditalic"))

    try:
        pdfmetrics.registerFont(TTFont(font_name, regular_path))
        # For bold/italic: use the real file if we have it; otherwise point the slot
        # at a built-in PostScript font that actually *looks* bold/italic. Without
        # this, <b> and <i> silently render as regular because the slot file IS the
        # regular file.
        bold_slot = f"{font_name}-Bold"
        italic_slot = f"{font_name}-Italic"
        bolditalic_slot = f"{font_name}-BoldItalic"
        if bold_path:
            pdfmetrics.registerFont(TTFont(bold_slot, bold_path))
        else:
            bold_slot = "Helvetica-Bold"
        if italic_path:
            pdfmetrics.registerFont(TTFont(italic_slot, italic_path))
        else:
            italic_slot = "Helvetica-Oblique"
        if bolditalic_path:
            pdfmetrics.registerFont(TTFont(bolditalic_slot, bolditalic_path))
        else:
            bolditalic_slot = "Helvetica-BoldOblique"
        registerFontFamily(
            font_name,
            normal=font_name,
            bold=bold_slot,
            italic=italic_slot,
            boldItalic=bolditalic_slot,
        )
        _REGISTERED_FONTS.add(font_name)
        return font_name
    except Exception:
        return "Helvetica"


def _flatten_contact(raw) -> str:
    if isinstance(raw, dict):
        return "  |  ".join(str(v) for v in raw.values() if v)
    if isinstance(raw, list):
        return "  |  ".join(str(v) for v in raw if v)
    return str(raw or "")


def _flatten_dates(d) -> str:
    if isinstance(d, list):
        return " ".join(str(x) for x in d)
    return str(d or "")


def _item_header_parts(item) -> tuple[str, str]:
    parts = []
    if item.get("title"): parts.append(str(item["title"]))
    if item.get("org"):   parts.append(str(item["org"]))
    return " &mdash; ".join(parts), _flatten_dates(item.get("dates", ""))


def _ed_parts(ed) -> tuple[str, str, str]:
    """Return (degree, school, dates) separately so the caller can bold only the degree."""
    return (
        str(ed.get("degree") or ""),
        str(ed.get("school") or ""),
        _flatten_dates(ed.get("dates", "")),
    )


def _skills_text(skills) -> str:
    if isinstance(skills, list):
        return " • ".join(str(s) for s in skills)
    if isinstance(skills, dict):
        return " • ".join(f"{k}: {v}" for k, v in skills.items())
    return str(skills or "")


def _summary_text(summary) -> str:
    if summary is None:
        return ""
    if isinstance(summary, str):
        return summary
    if isinstance(summary, list):
        return " ".join(_summary_text(s) for s in summary)
    if isinstance(summary, dict):
        # LLM sometimes wraps summary as {"text": "..."} or splits it into fields.
        for key in ("text", "summary", "content", "body", "value"):
            if key in summary and summary[key]:
                return _summary_text(summary[key])
        return " ".join(_summary_text(v) for v in summary.values() if v)
    return str(summary)


def build_pdf(resume: dict[str, Any], font_name: str, template_name: str = "",
               font_size: float = 10.0) -> bytes:
    effective_font = _register_pdf_font(font_name)
    style = _template_style(template_name)
    layout = style["layout"]
    accent = style["accent"]
    body_sz = max(8.0, min(13.0, float(font_size or 10.0)))
    head_sz = body_sz + 1
    name_sz = body_sz + 10
    contact_sz = body_sz - 0.5
    leading = body_sz + 3

    # Sidebar needs narrower outer margins because the inner table fills the width.
    is_sidebar = layout == "sidebar"
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.5 * inch if is_sidebar else 0.7 * inch,
        rightMargin=0.5 * inch if is_sidebar else 0.7 * inch,
        topMargin=0.4 * inch if layout == "banner" else 0.5 * inch,
        bottomMargin=0.5 * inch,
    )

    if layout == "shaded-header":
        heading_align = TA_LEFT
        heading_color = "#111111"
        bullet_glyph = "•"
        name_align = TA_CENTER
        name_case_upper = True
    elif layout == "serif-classic":
        heading_align = TA_CENTER
        heading_color = accent
        bullet_glyph = "—"
        name_align = TA_CENTER
        name_case_upper = True
    elif layout == "modern-bar":
        heading_align = TA_LEFT
        heading_color = accent
        bullet_glyph = "■"
        name_align = TA_LEFT
        name_case_upper = False
    elif layout == "banner":
        heading_align = TA_LEFT
        heading_color = accent
        bullet_glyph = "▸"
        name_align = TA_LEFT
        name_case_upper = False
    elif layout == "structured":
        heading_align = TA_LEFT
        heading_color = accent
        bullet_glyph = "•"
        name_align = TA_LEFT
        name_case_upper = True
    elif layout == "sidebar":
        heading_align = TA_LEFT
        heading_color = accent
        bullet_glyph = "•"
        name_align = TA_LEFT
        name_case_upper = False
    else:  # minimal
        heading_align = TA_CENTER
        heading_color = "#111"
        bullet_glyph = "•"
        name_align = TA_CENTER
        name_case_upper = False

    styles = {
        "name":        ParagraphStyle("name", fontName=effective_font, fontSize=name_sz,
                                       alignment=name_align, spaceAfter=2, leading=name_sz + 4,
                                       textColor="#111"),
        "name_banner": ParagraphStyle("name_banner", fontName=effective_font, fontSize=name_sz + 2,
                                       alignment=TA_LEFT, spaceAfter=2, leading=name_sz + 6,
                                       textColor="#ffffff"),
        "contact":     ParagraphStyle("contact", fontName=effective_font, fontSize=contact_sz,
                                       alignment=name_align, spaceAfter=8, leading=contact_sz + 2,
                                       textColor="#333"),
        "contact_banner": ParagraphStyle("contact_banner", fontName=effective_font, fontSize=contact_sz,
                                          alignment=TA_LEFT, spaceAfter=8, leading=contact_sz + 2,
                                          textColor="#f5f5f5"),
        "heading":     ParagraphStyle("heading", fontName=effective_font, fontSize=head_sz,
                                       textColor=heading_color, spaceBefore=8, spaceAfter=3,
                                       leading=head_sz + 2, alignment=heading_align),
        "body":        ParagraphStyle("body", fontName=effective_font, fontSize=body_sz,
                                       leading=leading, spaceAfter=2, alignment=TA_LEFT),
        "title":       ParagraphStyle("title", fontName=effective_font, fontSize=body_sz,
                                       leading=leading, spaceAfter=1, alignment=TA_LEFT),
        "sidebar_hd":  ParagraphStyle("sidebar_hd", fontName=effective_font, fontSize=body_sz,
                                       textColor="#ffffff", leading=leading, spaceBefore=6,
                                       spaceAfter=2),
        "sidebar_body": ParagraphStyle("sidebar_body", fontName=effective_font, fontSize=body_sz - 1,
                                        textColor="#ffffff", leading=leading - 1, spaceAfter=2),
    }

    name = resume.get("name") or "Your Name"
    if name_case_upper:
        name = name.upper()
    contact = _flatten_contact(resume.get("contact"))

    story = []

    # Usable page width (needed for shaded-header Table cells)
    _page_w = LETTER[0]
    _usable_w = _page_w - (0.5 * inch if is_sidebar else 0.7 * inch) * 2

    # Section-label overrides for RO Signature (match job-bot PDF exactly)
    _RO_SECTION_LABELS = {
        "summary":    "PROFESSIONAL SUMMARY",
        "skills":     "AREAS OF EXPERTISE",
        "experience": "PROFESSIONAL EXPERIENCE",
        "projects":   "PROJECTS",
        "education":  "EDUCATION",
        "certifications": "CERTIFICATIONS",
        "awards":     "CERTIFICATIONS & AWARDS",
    }

    # -------- header per layout --------
    if layout == "shaded-header":
        story.append(Paragraph(f"<b>{name}</b>", styles["name"]))
        headline = str(resume.get("headline") or resume.get("title") or "").strip()
        if headline:
            # Format subtitle with · dots between pipe-separated parts (matches job-bot PDF)
            subtitle_display = re.sub(r"\s*\|\s*", "  \u00b7  ", headline)
            story.append(Paragraph(
                f"<i>{subtitle_display}</i>",
                ParagraphStyle("ro_sub", fontName=effective_font, fontSize=body_sz + 1,
                               alignment=TA_CENTER, spaceAfter=3, leading=body_sz + 4,
                               textColor="#333333"),
            ))
        if contact:
            # Replace pipes with · for contact line too
            contact_display = re.sub(r"\s*\|\s*", "  \u00b7  ", contact)
            story.append(Paragraph(contact_display, styles["contact"]))
        story.append(HRFlowable(width="100%", thickness=1.0, color=HexColor("#111111"),
                                 spaceBefore=2, spaceAfter=6))
    elif layout == "banner":
        banner_tbl = Table(
            [[Paragraph(name, styles["name_banner"]),
              Paragraph(contact, styles["contact_banner"])]],
            colWidths=[3.5 * inch, 3.8 * inch],
        )
        banner_tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), HexColor(accent)),
            ("LEFTPADDING", (0, 0), (-1, -1), 16),
            ("RIGHTPADDING", (0, 0), (-1, -1), 16),
            ("TOPPADDING", (0, 0), (-1, -1), 14),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 14),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(banner_tbl)
        story.append(Spacer(1, 10))
    elif layout == "modern-bar":
        story.append(HRFlowable(width="100%", thickness=3, color=HexColor(accent),
                                 spaceBefore=0, spaceAfter=6))
        story.append(Paragraph(f'<font color="{accent}"><b>{name}</b></font>', styles["name"]))
        if contact:
            story.append(Paragraph(contact, styles["contact"]))
    elif layout == "serif-classic":
        story.append(Paragraph(f"<b>{name}</b>", styles["name"]))
        if contact:
            story.append(Paragraph(contact, styles["contact"]))
        story.append(HRFlowable(width="60%", thickness=0.6, color=HexColor(accent),
                                 hAlign="CENTER", spaceBefore=0, spaceAfter=4))
    elif layout == "structured":
        story.append(Paragraph(f"<b>{name}</b>", styles["name"]))
        if contact:
            story.append(Paragraph(contact, styles["contact"]))
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor(accent),
                                 spaceBefore=0, spaceAfter=6))
    elif layout == "sidebar":
        pass  # rendered as a table below, skip top header
    else:  # minimal
        story.append(Paragraph(f"<b>{name}</b>", styles["name"]))
        if contact:
            story.append(Paragraph(contact, styles["contact"]))

    # -------- sidebar layout: two-column table --------
    if layout == "sidebar":
        left_flow = [
            Paragraph(f"<b>{name}</b>",
                      ParagraphStyle("sn", fontName=effective_font, fontSize=15,
                                     textColor="#ffffff", leading=18, spaceAfter=6)),
        ]
        if contact:
            left_flow.append(Paragraph(contact, styles["sidebar_body"]))
        if resume.get("skills"):
            left_flow.append(Paragraph("<b>SKILLS</b>", styles["sidebar_hd"]))
            left_flow.append(Paragraph(_skills_text(resume["skills"]), styles["sidebar_body"]))
        if resume.get("education"):
            left_flow.append(Paragraph("<b>EDUCATION</b>", styles["sidebar_hd"]))
            for ed in resume["education"]:
                degree, school, d = _ed_parts(ed)
                tail = " &mdash; ".join(x for x in (school, d) if x)
                line = f"<b>{degree}</b>" + (f" &mdash; {tail}" if tail else "")
                left_flow.append(Paragraph(line, styles["sidebar_body"]))
        if resume.get("certifications"):
            left_flow.append(Paragraph("<b>CERTIFICATIONS</b>", styles["sidebar_hd"]))
            for c in resume["certifications"]:
                txt = c if isinstance(c, str) else c.get("name", str(c))
                left_flow.append(Paragraph(f"• {txt}", styles["sidebar_body"]))

        right_flow = []

        def r_heading(title):
            right_flow.append(Paragraph(f"<b>{title.upper()}</b>", styles["heading"]))

        if resume.get("summary"):
            r_heading("Summary")
            right_flow.append(Paragraph(_summary_text(resume["summary"]), styles["body"]))

        def r_section(title, items):
            if not items: return
            r_heading(title)
            for item in items:
                h, d = _item_header_parts(item)
                line = f"<b>{h}</b>   {d}" if d else f"<b>{h}</b>"
                right_flow.append(Paragraph(line, styles["title"]))
                for b in item.get("bullets", []) or []:
                    right_flow.append(Paragraph(f"{bullet_glyph} {str(b)}", styles["body"]))

        r_section("Experience", resume.get("experience") or [])
        r_section("Projects", resume.get("projects") or [])

        tbl = Table([[left_flow, right_flow]], colWidths=[2.3 * inch, 5.0 * inch])
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, 0), HexColor(accent)),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (0, 0), 12),
            ("RIGHTPADDING", (0, 0), (0, 0), 12),
            ("TOPPADDING", (0, 0), (0, 0), 14),
            ("BOTTOMPADDING", (0, 0), (0, 0), 14),
            ("LEFTPADDING", (1, 0), (1, 0), 14),
            ("RIGHTPADDING", (1, 0), (1, 0), 6),
            ("TOPPADDING", (1, 0), (1, 0), 8),
        ]))
        story.append(tbl)
        doc.build(story)
        return buf.getvalue()

    # -------- non-sidebar layouts: linear flow --------
    def heading_flowables(title):
        t = title.upper()
        if layout == "shaded-header":
            # Use the full human-readable label (e.g. "PROFESSIONAL EXPERIENCE")
            t = _RO_SECTION_LABELS.get(title.lower(), t)
            hdr_tbl = Table(
                [[Paragraph(f"<b>{t}</b>", styles["heading"])]],
                colWidths=[_usable_w],
            )
            hdr_tbl.setStyle(TableStyle([
                ("BACKGROUND",    (0, 0), (0, 0), HexColor("#e9e9e9")),
                ("LEFTPADDING",   (0, 0), (0, 0), 6),
                ("RIGHTPADDING",  (0, 0), (0, 0), 6),
                ("TOPPADDING",    (0, 0), (0, 0), 3),
                ("BOTTOMPADDING", (0, 0), (0, 0), 3),
            ]))
            return [Spacer(1, 4), hdr_tbl, Spacer(1, 3)]
        elif layout == "modern-bar":
            return [
                HRFlowable(width="100%", thickness=0.5, color=HexColor("#bbbbbb"), spaceBefore=6, spaceAfter=0),
                Paragraph(f'<font color="{accent}"><b>■ {t}</b></font>', styles["heading"]),
                HRFlowable(width="100%", thickness=0.5, color=HexColor(accent), spaceBefore=0, spaceAfter=4),
            ]
        elif layout == "banner":
            return [
                HRFlowable(width="100%", thickness=0.5, color=HexColor("#bbbbbb"), spaceBefore=6, spaceAfter=0),
                Paragraph(f'<font color="{accent}"><b>{t}</b></font>', styles["heading"]),
                HRFlowable(width="100%", thickness=0.5, color=HexColor("#cccccc"), spaceBefore=0, spaceAfter=4),
            ]
        elif layout == "serif-classic":
            return [
                HRFlowable(width="100%", thickness=0.75, color=HexColor("#222222"), spaceBefore=8, spaceAfter=0),
                Paragraph(f'<font color="{accent}"><b>{t}</b></font>', styles["heading"]),
                HRFlowable(width="100%", thickness=0.5, color=HexColor("#222222"), spaceBefore=0, spaceAfter=4),
            ]
        elif layout == "structured":
            return [
                HRFlowable(width="100%", thickness=0.5, color=HexColor("#bbbbbb"), spaceBefore=8, spaceAfter=0),
                Paragraph(f'<font color="{accent}"><b>{t}</b></font>', styles["heading"]),
                HRFlowable(width=1.2 * inch, thickness=1.2, color=HexColor(accent), hAlign="LEFT",
                           spaceBefore=0, spaceAfter=3),
            ]
        else:  # minimal — classic double-rule around the heading
            return [
                HRFlowable(width="100%", thickness=0.6, color=HexColor("#222222"), spaceBefore=8, spaceAfter=0),
                Paragraph(f"<b>{t}</b>", styles["heading"]),
                HRFlowable(width="100%", thickness=0.4, color=HexColor("#222222"), spaceBefore=0, spaceAfter=4),
            ]

    if resume.get("summary"):
        story.append(KeepTogether(
            heading_flowables("Summary") + [Paragraph(_summary_text(resume["summary"]), styles["body"])]
        ))

    if resume.get("skills"):
        story.append(KeepTogether(
            heading_flowables("Skills") + [Paragraph(_skills_text(resume["skills"]), styles["body"])]
        ))

    def _item_flowables(item):
        h, d = _item_header_parts(item)
        if layout == "shaded-header" and d:
            # Two-column table: bold title/org left, italic dates right — matches job-bot PDF
            _title_style = ParagraphStyle(
                "ro_role_left", fontName=effective_font, fontSize=body_sz,
                textColor="#111111", alignment=TA_LEFT, leading=leading,
                spaceBefore=3, spaceAfter=1,
            )
            _date_style = ParagraphStyle(
                "ro_role_right", fontName=effective_font, fontSize=body_sz - 0.5,
                textColor="#555555", alignment=TA_RIGHT, leading=leading,
                spaceBefore=3,
            )
            role_tbl = Table(
                [[Paragraph(f"<b>{h}</b>", _title_style),
                  Paragraph(f"<i>{d}</i>", _date_style)]],
                colWidths=[_usable_w * 0.68, _usable_w * 0.32],
            )
            role_tbl.setStyle(TableStyle([
                ("VALIGN",        (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING",   (0, 0), (-1, -1), 0),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 0),
                ("TOPPADDING",    (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
            ]))
            flows = [role_tbl]
        else:
            line = f"<b>{h}</b>   {d}" if d else f"<b>{h}</b>"
            flows = [Paragraph(line, styles["title"])]
        for b in item.get("bullets", []) or []:
            flows.append(Paragraph(f"{bullet_glyph} {str(b)}", styles["body"]))
        return flows

    def section(title, items):
        if not items:
            return
        # Keep heading together with at least the first item so the heading never orphans.
        first_item_flows = _item_flowables(items[0])
        story.append(KeepTogether(heading_flowables(title) + first_item_flows))
        for item in items[1:]:
            story.append(KeepTogether(_item_flowables(item)))

    section("Experience", resume.get("experience") or [])
    section("Projects", resume.get("projects") or [])

    if resume.get("education"):
        ed_flows = list(heading_flowables("Education"))
        for ed in resume["education"]:
            degree, school, d = _ed_parts(ed)
            tail = " &mdash; ".join(x for x in (school, d) if x)
            line = f"<b>{degree}</b>" + (f" &mdash; {tail}" if tail else "")
            ed_flows.append(Paragraph(line, styles["title"]))
            if ed.get("notes"):
                notes = ed["notes"]
                if isinstance(notes, list):
                    notes = " ".join(str(n) for n in notes)
                ed_flows.append(Paragraph(str(notes), styles["body"]))
        story.append(KeepTogether(ed_flows))

    if resume.get("certifications"):
        cert_flows = list(heading_flowables("Certifications"))
        for c in resume["certifications"]:
            txt = c if isinstance(c, str) else c.get("name", str(c))
            cert_flows.append(Paragraph(f"{bullet_glyph} {txt}", styles["body"]))
        story.append(KeepTogether(cert_flows))

    doc.build(story)
    return buf.getvalue()


# ------------------------- TOP-LEVEL -------------------------

def export_both_versions(resume: dict[str, Any], styled_font: str,
                          template_name: str = "", out_dir: str = "output",
                          font_size: float = 10.0):
    """Return dict of {'styled_docx': bytes, 'styled_pdf': bytes, 'ats_docx': bytes, 'ats_pdf': bytes}."""
    if styled_font not in FONTS:
        styled_font = "Lato"
    ats_font = fallback_for_ats(styled_font)

    return {
        "styled_docx": build_docx(resume, styled_font, font_size=font_size),
        "styled_pdf":  build_pdf(resume, styled_font, template_name, font_size=font_size),
        "ats_docx":    build_docx(resume, ats_font, font_size=font_size),
        "ats_pdf":     build_pdf(resume, ats_font, "", font_size=font_size),
        "styled_font": styled_font,
        "ats_font":    ats_font,
    }


# ------------------------- QR CODE HELPER -------------------------

def make_qr_png(url: str) -> bytes:
    """Return a PNG QR code for a URL (portfolio, LinkedIn). Used in header."""
    try:
        import qrcode
        img = qrcode.make(url)
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        return buf.getvalue()
    except Exception:
        return b""


# ------------------------- ATS PREFLIGHT -------------------------

def ats_preflight(pdf_bytes: bytes) -> dict:
    """Scan an exported PDF for common ATS risks."""
    from pypdf import PdfReader
    issues: list[str] = []
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        page_count = len(reader.pages)
        if page_count > 2:
            issues.append(f"{page_count} pages — most ATS prefer 1-2.")
        total_text = ""
        for p in reader.pages:
            try:
                total_text += p.extract_text() or ""
            except Exception:
                issues.append("A page failed text extraction — ATS may not parse it.")
        if len(total_text.strip()) < 400:
            issues.append("Very little extractable text — check for image-only content.")
        if "\t" in total_text:
            issues.append("Tabs found — some parsers split columns poorly.")
    except Exception as e:
        issues.append(f"PDF parse error: {e}")
        page_count = 0
        total_text = ""

    return {
        "page_count": page_count,
        "extractable_chars": len(total_text.strip()),
        "issues": issues,
        "pass": len(issues) == 0,
    }
