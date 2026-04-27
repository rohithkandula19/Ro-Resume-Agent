"""Parse resumes in many formats to plain text.

Supported:
  .pdf, .docx, .txt, .md, .rtf, .html, .htm, .png, .jpg, .jpeg (OCR when available)
"""

import io
import re

from pypdf import PdfReader
from docx import Document


def parse_pdf(file_bytes: bytes, max_pages: int = 5) -> str:
    """Extract text from the first `max_pages` pages. Resumes are rarely longer."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        if i >= max_pages:
            break
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n".join(pages).strip()


def parse_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts).strip()


def parse_rtf(file_bytes: bytes) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
        return rtf_to_text(file_bytes.decode("utf-8", errors="ignore")).strip()
    except ImportError:
        # Fallback: strip RTF control words with regex
        text = file_bytes.decode("utf-8", errors="ignore")
        text = re.sub(r"\\[a-z]+-?\d* ?", "", text)
        text = re.sub(r"[{}\\]", "", text)
        return text.strip()


def parse_html(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(text, "html.parser")
        for s in soup(["script", "style"]):
            s.decompose()
        return soup.get_text(separator="\n").strip()
    except ImportError:
        return re.sub(r"<[^>]+>", "\n", text).strip()


def parse_image(file_bytes: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(io.BytesIO(file_bytes))
        return pytesseract.image_to_string(img).strip()
    except ImportError:
        raise RuntimeError(
            "Image OCR needs pytesseract + Tesseract installed. "
            "Run: pip install pytesseract; brew install tesseract"
        )


def parse_resume(file_bytes: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    if name.endswith(".docx"):
        return parse_docx(file_bytes)
    if name.endswith(".rtf"):
        return parse_rtf(file_bytes)
    if name.endswith(".html") or name.endswith(".htm"):
        return parse_html(file_bytes)
    if name.endswith(".png") or name.endswith(".jpg") or name.endswith(".jpeg"):
        return parse_image(file_bytes)
    if name.endswith(".txt") or name.endswith(".md"):
        return file_bytes.decode("utf-8", errors="ignore").strip()
    # Last resort: try to decode as UTF-8 text
    try:
        text = file_bytes.decode("utf-8")
        if text.strip():
            return text.strip()
    except UnicodeDecodeError:
        pass
    raise ValueError(f"Unsupported file type: {filename}")
